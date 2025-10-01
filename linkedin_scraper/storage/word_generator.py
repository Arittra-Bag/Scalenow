"""Word document generator for knowledge repository data."""

from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    Document = None
    Inches = Pt = RGBColor = None
    WD_ALIGN_PARAGRAPH = WD_BREAK = WD_STYLE_TYPE = None
    OxmlElement = qn = None

from ..models.knowledge_item import KnowledgeItem, Category
from ..models.exceptions import StorageError
from ..utils.config import Config
from ..utils.logger import get_logger
from .repository_models import KnowledgeRepository

logger = get_logger(__name__)


class WordGenerator:
    """Generator for Word documents from knowledge repository data."""
    
    def __init__(self, config: Optional[Config] = None):
        """Initialize the Word generator."""
        self.config = config or Config.from_env()
        
        if Document is None:
            raise StorageError(
                "python-docx library not installed. "
                "Please install it with: pip install python-docx"
            )
        
        # Color scheme for categories
        self.category_colors = {
            Category.AI_MACHINE_LEARNING: RGBColor(52, 152, 219),  # Blue
            Category.SAAS_BUSINESS: RGBColor(241, 196, 15),        # Yellow
            Category.MARKETING_SALES: RGBColor(46, 204, 113),      # Green
            Category.LEADERSHIP_MANAGEMENT: RGBColor(155, 89, 182), # Purple
            Category.TECHNOLOGY_TRENDS: RGBColor(231, 76, 60),     # Red
            Category.COURSE_CONTENT: RGBColor(26, 188, 156),       # Teal
            Category.OTHER: RGBColor(149, 165, 166)                # Gray
        }
        
        logger.info("Word generator initialized")
    
    def generate_word_document(
        self,
        repository: KnowledgeRepository,
        output_path: str,
        include_toc: bool = True,
        include_metadata: bool = True,
        group_by_category: bool = True
    ) -> str:
        """Generate a comprehensive Word document from the knowledge repository."""
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            # Create document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add title page
            self._add_title_page(doc, repository)
            
            # Add table of contents placeholder
            if include_toc:
                self._add_table_of_contents_placeholder(doc)
            
            # Add metadata section
            if include_metadata:
                self._add_metadata_section(doc, repository)
            
            # Add executive summary
            self._add_executive_summary(doc, repository)
            
            # Add knowledge content
            if group_by_category:
                self._add_content_by_category(doc, repository)
            else:
                self._add_content_chronologically(doc, repository)
            
            # Add appendices
            self._add_course_references_appendix(doc, repository)
            self._add_source_links_appendix(doc, repository)
            
            # Save document
            doc.save(output_file)
            
            logger.info(f"Word document generated: {output_file}")
            return str(output_file)
            
        except Exception as e:
            raise StorageError(f"Failed to generate Word document: {e}")
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document."""
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = RGBColor(54, 96, 146)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading styles
        heading1 = styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(18)
        heading1.font.color.rgb = RGBColor(54, 96, 146)
        
        heading2 = styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(14)
        heading2.font.color.rgb = RGBColor(79, 129, 189)
        
        # Knowledge item style
        if 'Knowledge Item' not in [style.name for style in styles]:
            knowledge_style = styles.add_style('Knowledge Item', WD_STYLE_TYPE.PARAGRAPH)
            knowledge_font = knowledge_style.font
            knowledge_font.name = 'Calibri'
            knowledge_font.size = Pt(11)
            knowledge_style.paragraph_format.space_after = Pt(6)
            knowledge_style.paragraph_format.left_indent = Inches(0.25)
        
        # Source link style
        if 'Source Link' not in [style.name for style in styles]:
            source_style = styles.add_style('Source Link', WD_STYLE_TYPE.PARAGRAPH)
            source_font = source_style.font
            source_font.name = 'Calibri'
            source_font.size = Pt(9)
            source_font.color.rgb = RGBColor(79, 129, 189)
            source_font.italic = True
            source_style.paragraph_format.space_after = Pt(12)
            source_style.paragraph_format.left_indent = Inches(0.25)
    
    def _add_title_page(self, doc: Document, repository: KnowledgeRepository):
        """Add a title page to the document."""
        # Main title
        title = doc.add_paragraph('LinkedIn Knowledge Repository', style='Custom Title')
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Curated Insights and Knowledge Base')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(89, 89, 89)
        
        # Add some space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Repository stats
        stats_para = doc.add_paragraph()
        stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        stats_text = f"""
        Total Knowledge Items: {len(repository.items)}
        Categories Covered: {len(set(item.category for item in repository.items))}
        Generated: {datetime.now().strftime('%B %d, %Y')}
        Version: {repository.version}
        """
        
        run = stats_para.add_run(stats_text.strip())
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(89, 89, 89)
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents_placeholder(self, doc: Document):
        """Add a table of contents placeholder."""
        doc.add_heading('Table of Contents', level=1)
        
        toc_para = doc.add_paragraph()
        run = toc_para.add_run('[Table of Contents will be generated when document is opened in Microsoft Word]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(89, 89, 89)
        
        doc.add_page_break()
    
    def _add_metadata_section(self, doc: Document, repository: KnowledgeRepository):
        """Add a metadata section with repository information."""
        doc.add_heading('Repository Information', level=1)
        
        # Basic information
        info_items = [
            ('Repository Version', repository.version),
            ('Created', repository.created_at.strftime('%Y-%m-%d %H:%M:%S')),
            ('Last Updated', repository.updated_at.strftime('%Y-%m-%d %H:%M:%S')),
            ('Total Items', str(len(repository.items))),
        ]
        
        for label, value in info_items:
            para = doc.add_paragraph()
            run = para.add_run(f"{label}: ")
            run.bold = True
            para.add_run(value)
        
        # Category distribution
        if repository.items:
            doc.add_heading('Category Distribution', level=2)
            
            category_counts = {}
            for item in repository.items:
                category = item.category.value
                category_counts[category] = category_counts.get(category, 0) + 1
            
            for category, count in sorted(category_counts.items(), key=lambda x: x[1], reverse=True):
                percentage = (count / len(repository.items)) * 100
                para = doc.add_paragraph()
                run = para.add_run(f"• {category}: ")
                run.bold = True
                para.add_run(f"{count} items ({percentage:.1f}%)")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, repository: KnowledgeRepository):
        """Add an executive summary section."""
        doc.add_heading('Executive Summary', level=1)
        
        if not repository.items:
            doc.add_paragraph("No knowledge items found in the repository.")
            return
        
        # Summary statistics
        total_items = len(repository.items)
        categories = len(set(item.category for item in repository.items))
        
        # Recent activity
        recent_items = sorted(repository.items, key=lambda x: x.extraction_date, reverse=True)[:5]
        
        # Top topics
        topic_counts = {}
        for item in repository.items:
            topic_counts[item.topic] = topic_counts.get(item.topic, 0) + 1
        top_topics = sorted(topic_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        
        # Course references
        total_courses = sum(len(item.course_references) for item in repository.items)
        
        summary_text = f"""
        This knowledge repository contains {total_items} curated insights across {categories} different categories, 
        providing a comprehensive resource for business intelligence and learning.
        
        The repository includes {total_courses} course and training references, making it a valuable resource 
        for professional development and skill enhancement.
        """
        
        doc.add_paragraph(summary_text.strip())
        
        # Key highlights
        doc.add_heading('Key Highlights', level=2)
        
        highlights = [
            f"Most active category: {max(set(item.category.value for item in repository.items), key=lambda x: sum(1 for item in repository.items if item.category.value == x))}",
            f"Top topic: {top_topics[0][0] if top_topics else 'N/A'}",
            f"Latest addition: {recent_items[0].post_title[:50]}..." if recent_items else "N/A",
            f"Course references: {total_courses} training opportunities identified"
        ]
        
        for highlight in highlights:
            para = doc.add_paragraph()
            para.add_run(f"• {highlight}")
        
        doc.add_page_break()
    
    def _add_content_by_category(self, doc: Document, repository: KnowledgeRepository):
        """Add knowledge content organized by category."""
        doc.add_heading('Knowledge Content by Category', level=1)
        
        # Group items by category
        category_items = {}
        for item in repository.items:
            category = item.category
            if category not in category_items:
                category_items[category] = []
            category_items[category].append(item)
        
        # Sort categories by item count (descending)
        sorted_categories = sorted(category_items.items(), key=lambda x: len(x[1]), reverse=True)
        
        for category, items in sorted_categories:
            # Category header
            category_heading = doc.add_heading(category.value, level=1)
            
            # Color the heading
            if category in self.category_colors:
                for run in category_heading.runs:
                    run.font.color.rgb = self.category_colors[category]
            
            # Category summary
            doc.add_paragraph(f"Total items in this category: {len(items)}")
            
            # Sort items by extraction date (newest first)
            sorted_items = sorted(items, key=lambda x: x.extraction_date, reverse=True)
            
            for idx, item in enumerate(sorted_items, 1):
                self._add_knowledge_item(doc, item, idx)
            
            # Add page break between categories (except for the last one)
            if category != sorted_categories[-1][0]:
                doc.add_page_break()
    
    def _add_content_chronologically(self, doc: Document, repository: KnowledgeRepository):
        """Add knowledge content in chronological order."""
        doc.add_heading('Knowledge Content (Chronological)', level=1)
        
        # Sort items by extraction date (newest first)
        sorted_items = sorted(repository.items, key=lambda x: x.extraction_date, reverse=True)
        
        current_month = None
        for idx, item in enumerate(sorted_items, 1):
            # Add month separator
            item_month = item.extraction_date.strftime('%Y-%m') if item.extraction_date else 'Unknown'
            if item_month != current_month:
                if current_month is not None:
                    doc.add_paragraph()  # Add space
                
                month_heading = doc.add_heading(
                    item.extraction_date.strftime('%B %Y') if item.extraction_date else 'Unknown Date',
                    level=2
                )
                current_month = item_month
            
            self._add_knowledge_item(doc, item, idx)
    
    def _add_knowledge_item(self, doc: Document, item: KnowledgeItem, index: int):
        """Add a single knowledge item to the document."""
        # Item header
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(f"{index}. {item.post_title}")
        header_run.bold = True
        header_run.font.size = Pt(12)
        
        # Category and topic
        meta_para = doc.add_paragraph()
        category_run = meta_para.add_run(f"Category: {item.category.value}")
        category_run.font.color.rgb = self.category_colors.get(item.category, RGBColor(89, 89, 89))
        category_run.bold = True
        
        meta_para.add_run(f" | Topic: {item.topic}")
        
        if item.extraction_date:
            meta_para.add_run(f" | Date: {item.extraction_date.strftime('%Y-%m-%d')}")
        
        # Knowledge content
        content_para = doc.add_paragraph(style='Knowledge Item')
        content_para.add_run(item.key_knowledge_content)
        
        # Infographic summary (if available)
        if item.infographic_summary:
            infographic_para = doc.add_paragraph()
            infographic_run = infographic_para.add_run("Visual Content: ")
            infographic_run.bold = True
            infographic_run.italic = True
            infographic_para.add_run(item.infographic_summary)
        
        # Course references (if available)
        if item.course_references:
            course_para = doc.add_paragraph()
            course_run = course_para.add_run("Related Courses: ")
            course_run.bold = True
            course_run.font.color.rgb = RGBColor(26, 188, 156)  # Teal
            course_para.add_run(', '.join(item.course_references))
        
        # Notes/Applications (if available)
        if item.notes_applications:
            notes_para = doc.add_paragraph()
            notes_run = notes_para.add_run("Key Takeaways: ")
            notes_run.bold = True
            notes_run.font.color.rgb = RGBColor(155, 89, 182)  # Purple
            notes_para.add_run(item.notes_applications)
        
        # Source link
        source_para = doc.add_paragraph(style='Source Link')
        source_run = source_para.add_run(f"Source: {item.source_link}")
        
        # Add separator line
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("─" * 80)
        separator_run.font.color.rgb = RGBColor(200, 200, 200)
        separator_para.add_run("\n")
    
    def _add_course_references_appendix(self, doc: Document, repository: KnowledgeRepository):
        """Add an appendix with all course references."""
        # Collect all course references
        all_courses = []
        for item in repository.items:
            if item.course_references:
                for course in item.course_references:
                    all_courses.append({
                        'course': course,
                        'category': item.category.value,
                        'topic': item.topic,
                        'source': item.source_link
                    })
        
        if not all_courses:
            return
        
        doc.add_page_break()
        doc.add_heading('Appendix A: Course References', level=1)
        
        # Group by category
        course_by_category = {}
        for course_info in all_courses:
            category = course_info['category']
            if category not in course_by_category:
                course_by_category[category] = []
            course_by_category[category].append(course_info)
        
        for category, courses in sorted(course_by_category.items()):
            doc.add_heading(category, level=2)
            
            # Remove duplicates
            unique_courses = {}
            for course_info in courses:
                course_name = course_info['course']
                if course_name not in unique_courses:
                    unique_courses[course_name] = course_info
            
            for course_name, course_info in sorted(unique_courses.items()):
                para = doc.add_paragraph()
                course_run = para.add_run(f"• {course_name}")
                course_run.bold = True
                
                para.add_run(f"\n  Topic: {course_info['topic']}")
                para.add_run(f"\n  Source: {course_info['source']}")
                para.add_run("\n")
    
    def _add_source_links_appendix(self, doc: Document, repository: KnowledgeRepository):
        """Add an appendix with all source links."""
        doc.add_page_break()
        doc.add_heading('Appendix B: Source Links', level=1)
        
        # Group by category
        links_by_category = {}
        for item in repository.items:
            category = item.category.value
            if category not in links_by_category:
                links_by_category[category] = []
            
            links_by_category[category].append({
                'title': item.post_title,
                'link': item.source_link,
                'date': item.extraction_date.strftime('%Y-%m-%d') if item.extraction_date else 'Unknown'
            })
        
        for category, links in sorted(links_by_category.items()):
            doc.add_heading(category, level=2)
            
            # Sort by date (newest first)
            sorted_links = sorted(links, key=lambda x: x['date'], reverse=True)
            
            for link_info in sorted_links:
                para = doc.add_paragraph()
                title_run = para.add_run(f"• {link_info['title']}")
                title_run.bold = True
                
                para.add_run(f"\n  Date: {link_info['date']}")
                para.add_run(f"\n  URL: {link_info['link']}")
                para.add_run("\n")
    
    def generate_summary_document(
        self,
        repository: KnowledgeRepository,
        output_path: str,
        max_items_per_category: int = 5
    ) -> str:
        """Generate a summary document with top items from each category."""
        try:
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)
            
            doc = Document()
            self._setup_document_styles(doc)
            
            # Title
            doc.add_paragraph('LinkedIn Knowledge Repository - Executive Summary', style='Custom Title')
            
            # Summary stats
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph(f"Total Items: {len(repository.items)}")
            doc.add_paragraph()
            
            # Group by category and get top items
            category_items = {}
            for item in repository.items:
                category = item.category
                if category not in category_items:
                    category_items[category] = []
                category_items[category].append(item)
            
            # Sort categories by item count
            sorted_categories = sorted(category_items.items(), key=lambda x: len(x[1]), reverse=True)
            
            for category, items in sorted_categories:
                doc.add_heading(f"{category.value} ({len(items)} items)", level=1)
                
                # Sort items by extraction date and take top N
                top_items = sorted(items, key=lambda x: x.extraction_date, reverse=True)[:max_items_per_category]
                
                for idx, item in enumerate(top_items, 1):
                    # Simplified item format
                    para = doc.add_paragraph()
                    title_run = para.add_run(f"{idx}. {item.post_title}")
                    title_run.bold = True
                    
                    # Shortened content
                    content = item.key_knowledge_content
                    if len(content) > 200:
                        content = content[:200] + "..."
                    
                    doc.add_paragraph(content, style='Knowledge Item')
                    
                    # Source
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run(f"Source: {item.source_link}")
                    source_run.font.size = Pt(9)
                    source_run.italic = True
                    
                    doc.add_paragraph()  # Space between items
            
            doc.save(output_file)
            
            logger.info(f"Summary Word document generated: {output_file}")
            return str(output_file)
            
        except Exception as e:
            raise StorageError(f"Failed to generate summary document: {e}")
    
    def update_existing_document(
        self,
        doc_path: str,
        new_items: List[KnowledgeItem]
    ) -> str:
        """Update an existing Word document with new knowledge items."""
        try:
            doc_file = Path(doc_path)
            if not doc_file.exists():
                raise StorageError(f"Word document not found: {doc_path}")
            
            # Load existing document
            doc = Document(doc_path)
            
            # Add new section for updates
            doc.add_page_break()
            doc.add_heading(f'Recent Updates - {datetime.now().strftime("%B %d, %Y")}', level=1)
            
            # Add new items
            for idx, item in enumerate(new_items, 1):
                self._add_knowledge_item(doc, item, idx)
            
            # Save updated document
            doc.save(doc_path)
            
            logger.info(f"Word document updated with {len(new_items)} new items: {doc_path}")
            return str(doc_path)
            
        except Exception as e:
            raise StorageError(f"Failed to update Word document: {e}")
    
    def get_next_version_filename(self, base_path: str) -> str:
        """Get the next version filename for Word documents."""
        base_path = Path(base_path)
        base_name = base_path.stem
        extension = base_path.suffix
        directory = base_path.parent
        
        # Find existing versions
        version_pattern = re.compile(rf"{re.escape(base_name)}_v(\d+){re.escape(extension)}")
        versions = []
        
        for file in directory.glob(f"{base_name}_v*{extension}"):
            match = version_pattern.match(file.name)
            if match:
                versions.append(int(match.group(1)))
        
        # Get next version number
        next_version = max(versions, default=0) + 1
        
        return str(directory / f"{base_name}_v{next_version}{extension}")